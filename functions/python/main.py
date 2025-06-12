from firebase_functions import firestore_fn, logger, params
import firebase_admin
from firebase_admin import initialize_app, storage, credentials
from google.auth.transport.requests import Request
from google.oauth2 import id_token
import tempfile
import os
from docx import Document
import re
import requests
from datetime import datetime, timezone
import json

def initialize_firebase(app_name):
    try:
        service_account_json = os.environ.get("DOCGEN_SA")
        if not service_account_json:
            print("❌ DOCGEN_SA not set in Firebase config!")
            raise ValueError("❌ DOCGEN_SA not set in Firebase config!")
        
        service_account_info = json.loads(service_account_json)
        cred = credentials.Certificate(service_account_info)
        
        return initialize_app(cred, name=app_name)
    
    except Exception as e:
        print(f"🔥 Firebase init failed: {str(e)}")
        raise

# Triggered when a new template document is created
@firestore_fn.on_document_created(document="templates/{templateId}")
def processtemplate(event: firestore_fn.Event[firestore_fn.DocumentSnapshot]) -> None:
    """Extracts placeholders from uploaded template and updates document."""
    
    snapshot = event.data
    if not snapshot.exists:
        print("Document no longer exists")
        return

    template_data = snapshot.to_dict()
    template_id = event.params["templateId"]

    try:
        snapshot.reference.update({"status": "processing"})

        file_path = template_data["storagePath"]
        bucket_name = template_data["downloadURL"].split('/')[5]
        bucket = storage.bucket(bucket_name)
        
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
            temp_path = temp_file.name
            blob = bucket.blob(file_path)
            blob.download_to_filename(temp_path)
            print(f"Downloaded template to {temp_path}")

            placeholders = extract_placeholders(temp_path)
            print(f"Found placeholders: {placeholders}")

            placeholder_config = {
                p: {
                    "type": "long_text",
                    "required": True,
                    "alias": p.replace('_', ' ').title()
                } for p in placeholders
            }

            snapshot.reference.update({
                "placeholders": placeholder_config,
                "status": "processed"
            })
            print("Template processed successfully")

    except Exception as error:
        print(f"Error processing template: {error}")
        snapshot.reference.update({
            "status": "failed",
            "error": str(error)
        })
        raise

    finally:
        # Clean up temporary file
        if 'temp_path' in locals() and os.path.exists(temp_path):
            os.remove(temp_path)

def extract_placeholders(docx_path: str) -> list:
    """Extracts all placeholder patterns from a DOCX file."""
    doc = Document(docx_path)
    placeholders = set()

    # Check paragraphs
    for paragraph in doc.paragraphs:
        matches = re.findall(r'\{\{([^}]+)\}\}', paragraph.text)
        placeholders.update(matches)

    # Check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                matches = re.findall(r'\{\{([^}]+)\}\}', cell.text)
                placeholders.update(matches)

    return list(placeholders)

# This function is triggered when a new document is created in the document_jobs collection
# It calls the Cloud Run service to process the document
@firestore_fn.on_document_created(document="document_jobs/{docId}")
def process_document_job(event: firestore_fn.Event[firestore_fn.DocumentSnapshot]) -> None:
    snapshot = event.data
    if not snapshot.exists:
        print("Document no longer exists")
        return
    
    app_name = 'process_document_job_app'
    if firebase_admin._apps.get(app_name) is None:
        initialize_firebase(app_name)

    try:
        doc_data = snapshot.to_dict()
        job_id = event.params["docId"]
        form_id = doc_data.get("formId")
        
        if not form_id:
            print("No formId found in document")
            return

        # Update job status to 'pickedup'
        snapshot.reference.update({
            "status": "pickedup",
            "updatedAt": datetime.now(timezone.utc)
        })

        cloud_run_url = os.environ.get("DOCGEN_URL")
        if not cloud_run_url:
            raise ValueError("DOCGEN_URL not set in Firebase config")
        
        # Ensure URL ends with / if not present
        cloud_run_url = cloud_run_url.rstrip('/') + '/'
        
        headers = {
            "Authorization": f"Bearer {get_cloud_run_token(cloud_run_url)}",
            "Content-Type": "application/json"
        }
        
        response = requests.post(
            cloud_run_url,
            headers=headers,
            json={
                "jobId": job_id,
                "formId": form_id
            },
            timeout=30
        )

        if response.status_code != 200:
            error_msg = f"Cloud Run request failed: {response.status_code} - {response.text}"
            print(error_msg)
            snapshot.reference.update({
                "status": "failed",
                "error": error_msg[:500],
                "updatedAt": datetime.now(datetime.timezone.utc)
            })
        else:
            print(f"Successfully processed job {job_id} for form {form_id}")

    except Exception as error:
        error_msg = f"Error processing document job: {error}"
        print(error_msg)
        if snapshot.exists:
            snapshot.reference.update({
                "status": "failed",
                "error": str(error)[:500],
                "updatedAt": datetime.now(datetime.timezone.utc)
            })
        raise

def get_cloud_run_token(target_audience):
    try:
        return id_token.fetch_id_token(Request(), target_audience)
    except Exception as e:
        print(f"Error getting ID token: {str(e)}")
        raise
