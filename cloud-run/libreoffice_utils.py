import subprocess
import os
import logging
from pathlib import Path
import shutil
from docx import Document
import tempfile

logger = logging.getLogger(__name__)

def fill_template_and_convert(template_path: str, form_data: dict) -> tuple[str, str]:
    """
    Fills a Word template with form data and converts it to both DOCX and PDF formats.
    
    Args:
        template_path: Path to the Word template file
        form_data: Dictionary containing form data to fill in the template
        
    Returns:
        Tuple of (filled_docx_path, pdf_path)
    """
    try:
        # Validate input file
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template file not found: {template_path}")

        # Create a temporary directory for output files
        temp_dir = tempfile.mkdtemp()
        base_name = os.path.splitext(os.path.basename(template_path))[0]
        filled_docx_path = os.path.join(temp_dir, f"{base_name}_filled.docx")
        
        # Fill the template with form data
        fill_word_template(template_path, filled_docx_path, form_data)
        
        # Convert to PDF
        pdf_path = convert_to_pdf(filled_docx_path)
        
        return filled_docx_path, pdf_path
        
    except Exception as e:
        print(f"Template processing failed: {str(e)}")
        raise

def fill_word_template(template_path: str, output_path: str, form_data: dict):
    """
    Fills a Word template with form data by replacing parameter placeholders.
    
    Args:
        template_path: Path to the Word template file
        output_path: Path to save the filled document
        form_data: Dictionary containing form data to fill in the template
    """
    try:
        doc = Document(template_path)
        
        # Replace parameters in paragraphs
        for paragraph in doc.paragraphs:
            for key, field_data in form_data.items():
                # Extract the actual value from the field data object
                value = str(field_data.get('value', '')) if isinstance(field_data, dict) else str(field_data)
                placeholder = f"{{{{{key}}}}}"
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, value)
        
        # Replace parameters in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, field_data in form_data.items():
                            # Extract the actual value from the field data object
                            value = str(field_data.get('value', '')) if isinstance(field_data, dict) else str(field_data)
                            placeholder = f"{{{{{key}}}}}"
                            if placeholder in paragraph.text:
                                paragraph.text = paragraph.text.replace(placeholder, value)

        doc.save(output_path)
        print(f"Successfully filled template: {output_path}")
        
    except Exception as e:
        print(f"Failed to fill Word template: {str(e)}")
        raise

def convert_to_pdf(docx_path: str) -> str:
    try:
        # Validate input file
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"Input file not found: {docx_path}")

        # Find LibreOffice binary
        libreoffice_path = shutil.which("libreoffice") or shutil.which("soffice")
        if not libreoffice_path:
            raise RuntimeError("LibreOffice not found in PATH")

        # Prepare output directory
        output_dir = os.path.dirname(docx_path)
        output_pdf = os.path.splitext(docx_path)[0] + ".pdf"

        # Run LibreOffice in headless mode
        cmd = [
            libreoffice_path,
            "--headless",
            "--convert-to", "pdf",
            "--outdir", output_dir,
            docx_path
        ]
        
        print(f"Running command: {' '.join(cmd)}")
        result = subprocess.run(
            cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            check=True
        )
        
        # Log LibreOffice output
        if result.stdout:
            print(f"LibreOffice stdout: {result.stdout}")
        if result.stderr:
            print(f"LibreOffice stderr: {result.stderr}")

        if not os.path.exists(output_pdf):
            raise RuntimeError(f"PDF not generated at {output_pdf}")

        print(f"Successfully converted to PDF: {output_pdf}")
        return output_pdf

    except subprocess.CalledProcessError as e:
        print(f"LibreOffice failed with code {e.returncode}: {e.stderr}")
        raise
    except Exception as e:
        print(f"PDF conversion failed: {str(e)}")
        raise